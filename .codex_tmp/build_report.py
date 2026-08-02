from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Admin\Documents\GitHub\AI-Tools-Project")
OUT = ROOT / "Final_Submission" / "Egyptian_ALPR_Final_Report.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "2E74B5"
PALE = "EAF2F8"
GRAY = "5B6573"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

def set_font(run, name="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Title", 30, NAVY, 0, 8), ("Subtitle", 14, GRAY, 0, 10),
    ("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4)]:
    s = styles[name]
    s.font.name = "Calibri"; s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW") or OxmlElement("w:tblW")
    if tblW.getparent() is None: tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd") or OxmlElement("w:tblInd")
    if tblInd.getparent() is None: tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c, PALE)
        for r in c.paragraphs[0].runs: set_font(r, size=10, bold=True, color=NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text=str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: set_font(r, size=9.5)
    set_table_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def add_caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    set_font(p.add_run(text), size=9, italic=True, color=GRAY)

def add_figure(path, caption, width=6.2):
    if path.exists():
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
        p.add_run().add_picture(str(path), width=Inches(width)); add_caption(caption)

def page_break(): doc.add_page_break()

def add_bullets(items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(4); p.add_run(item)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(115); p.paragraph_format.space_after=Pt(16)
set_font(p.add_run("FINAL PROJECT REPORT"), size=12, bold=True, color=BLUE)
p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Egyptian Automatic License Plate Recognition Using a Multi-Stage YOLO Pipeline")
p=doc.add_paragraph(style="Subtitle"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Dataset engineering, plate localization, Arabic character detection, and an interactive deployment workbench")
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(85); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Student Name: ______________________________\nStudent ID: __________________________________\nSupervisor: __________________________________\nDepartment / University: _______________________\nSubmission Date: ______________________________"), size=11, color=GRAY)

page_break()
doc.add_heading("Abstract", level=1)
doc.add_paragraph("This project develops an automatic license plate recognition (ALPR) pipeline specialized for Egyptian vehicle and motorcycle plates. The work combines a reproducible seven-stage dataset-engineering workflow with a three-stage vision cascade: pretrained vehicle detection, custom license-plate detection, and Arabic digit/letter detection with right-to-left decoding. Two raw plate datasets containing 2,551 images were inspected; 2,317 valid images were harmonized into a unified YOLO representation and split into 1,622 training, 348 validation, and 347 test images. The selected YOLO11 plate detector reached 96.24% mAP@0.50 and 82.38% mAP@0.50:0.95 on validation at epoch 49. The YOLO26 character detector reached 97.93% mAP@0.50 and 63.75% mAP@0.50:0.95 at its best recorded epoch. These results demonstrate strong component-level detection, while final plate-string accuracy, end-to-end recall, latency, and identity-grouped test performance remain to be measured before deployment claims can be made.")
doc.add_heading("Keywords", level=2)
doc.add_paragraph("Automatic license plate recognition; Egyptian license plates; Arabic OCR; YOLO; object detection; dataset harmonization; intelligent transportation systems.")

doc.add_heading("Table of Contents", level=1)
for item in ["Chapter 1 - Introduction", "1.1 Problem Statement", "1.2 Objectives and Scope", "1.3 Report Organization", "Chapter 2 - Background and Literature Review", "Chapter 3 - Dataset, Proposed Method, Evaluation, and Results", "Chapter 4 - Conclusions and Future Work", "References"]:
    doc.add_paragraph(item)

page_break(); doc.add_heading("Chapter 1: Introduction", level=1)
doc.add_paragraph("Automatic license plate recognition converts road-scene imagery into structured vehicle identifiers. It supports parking access, traffic analytics, tolling, and law-enforcement workflows, but performance depends on the visual domain. Egyptian plates combine Arabic letters and digits, multiple plate colors and layouts, small character regions, right-to-left reading order, and imagery affected by distance, blur, glare, occlusion, and perspective. A practical system must therefore address data quality, plate localization, character recognition, output formatting, and deployment as one traceable workflow.")
doc.add_paragraph("The project implements that workflow as reusable Python modules, command-line tools, a FastAPI backend, and a React/Vite interface. It also preserves training artifacts and dataset reports so that model performance can be traced to the data and configuration that produced it.")
doc.add_heading("1.1 Problem Statement", level=2)
doc.add_paragraph("The central problem is to detect and recognize Egyptian vehicle and motorcycle license plates reliably from unconstrained images and video. General-purpose OCR is poorly matched to tiny Arabic glyphs, nonuniform plate crops, and domain-specific ordering rules. Meanwhile, raw datasets arrive in different annotation formats and contain missing labels, near duplicates, corrupted samples, and inconsistent resolutions. The project therefore asks: how can heterogeneous Egyptian plate datasets be transformed into a reproducible training resource, and how effectively can a staged detector localize plates and their Arabic characters?")
doc.add_heading("1.2 Objectives and Scope", level=2)
add_bullets(["Inspect, validate, harmonize, preprocess, and split heterogeneous plate datasets reproducibly.", "Train a one-class YOLO11 detector for Egyptian license plates.", "Train a 38-class YOLO26 detector for Arabic letters and digits on plate crops.", "Order detections by row and right-to-left position, then map transliterated labels to Arabic glyphs.", "Expose training and inference through a FastAPI service and React workbench.", "Report component-level results honestly and reserve end-to-end recognition claims for a leakage-controlled test benchmark."])
doc.add_heading("1.3 Report Organization", level=2)
doc.add_paragraph("Chapter 2 introduces ALPR concepts and reviews classical, CNN-based, YOLO-based, and sequence-recognition approaches. Chapter 3 describes the datasets, preprocessing, proposed cascade, evaluation metrics, and results obtained to date. Chapter 4 summarizes the contribution, limitations, and prioritized future work. References provide the scholarly and software sources used in the design and interpretation.")

page_break(); doc.add_heading("Chapter 2: Background", level=1)
doc.add_heading("2.1 ALPR Pipeline", level=2)
doc.add_paragraph("ALPR systems usually contain plate localization, geometric/photometric normalization, character recognition, and syntactic post-processing. Classical systems used edges, morphology, connected components, and template matching. Modern systems increasingly use convolutional detectors because they learn discriminative features and handle broader variation in pose, illumination, and background. Recognition can be performed as segmented character detection or as sequence transcription using models such as CRNN with connectionist temporal classification.")
doc.add_heading("2.2 Why a Cascaded Detector", level=2)
doc.add_paragraph("A cascade reduces the search area at each stage. A vehicle detector first identifies relevant road objects; plate detection then operates on padded vehicle crops; character detection operates on the much smaller plate crop. This decomposition matches the scale hierarchy of the problem and enables separate diagnosis of vehicle, plate, and character errors. Its cost is error propagation: a missed vehicle prevents all downstream recognition, and an inaccurate plate crop reduces character quality.")
doc.add_heading("2.3 Egyptian Plate Characteristics", level=2)
doc.add_paragraph("Egyptian plates require locale-aware decoding. Letters and digits occupy distinct regions, and detected symbols must be ordered by row and then from right to left. Motorcycles introduce additional aspect ratios and viewing angles. Dataset diversity is therefore as important as the choice of detector.")

doc.add_heading("2.4 Literature Review", level=1)
doc.add_paragraph("Early Egyptian ALPR research used plate-region extraction, segmentation, and template matching. El-Adawi et al. described a three-part system for Egyptian plates and emphasized video processing without additional sensors [1]. Abd El Rahman et al. combined localization and skew correction, connected-component segmentation, and adapted template matching, reporting 81% accuracy on a two-hour video [2]. These approaches are interpretable but depend heavily on thresholds and segmentation quality.")
doc.add_paragraph("Deep learning shifted the task toward learned localization and recognition. Redmon et al. formulated object detection as a single regression problem in YOLO, enabling real-time inference [3]. Masood et al. demonstrated a cascade of convolutional networks for license-plate detection and recognition under variations in pose, lighting, occlusion, and plate templates [4]. Shi et al.'s CRNN integrates feature extraction, sequence modeling, and transcription, making it a useful optional fallback where character segmentation is unreliable [5].")
doc.add_paragraph("For the Egyptian domain, Youssef et al. introduced the EALPR benchmark and a Tiny-YOLOv3 system, reporting mAP values of 97.89% for plate detection and 92.46% for character recognition [6], [7]. Abdellatif et al. reported 93% recognition accuracy on 200 Egyptian plate images using edge-based localization, segmentation, and recognition [8]. Recent work continues to emphasize the shortage of diverse Egyptian motorcycle and vehicle data and the need for robust evaluation across plate types [9].")
doc.add_paragraph("The present project differs by treating dataset engineering, plate and character training, evaluation, API integration, and human review as one reproducible system. However, its current metrics are component-level validation metrics and should not be compared directly with whole-string accuracy reported by other studies.")

page_break(); doc.add_heading("Chapter 3: Data Set Description", level=1)
doc.add_paragraph("The plate-localization corpus combines two raw sources with different annotation formats. Dataset A uses Pascal VOC XML and contains 464 JPEG images; Dataset B uses YOLO text labels and contains 2,087 images. Automated inspection found 2,551 images and 2,550 annotations in total. Harmonization converted 2,317 usable samples to JPG images with normalized YOLO labels.")
add_table(["Dataset", "Images", "Annotation format", "Average resolution", "Key quality signals"], [
    ["A", "464", "Pascal VOC XML", "1849 x 2421", "2 images without annotations; 231 near-duplicate pairs"],
    ["B", "2,087", "YOLO TXT", "883 x 904", "1 corrupted image; 4 exact-duplicate groups; 1 out-of-bounds box"],
    ["Unified", "2,317", "YOLO normalized TXT", "Mixed", "Stratified 70/15/15 split"]], [900,900,1500,1600,4460])
add_figure(ROOT/"reports/figures/dataset_size_comparison.png", "Figure 3.1. Raw dataset image counts generated by the project EDA pipeline.", 6.0)
doc.add_heading("3.1 Plate Detection Split", level=2)
add_table(["Split", "Images", "Percentage"], [["Training","1,622","70.0%"],["Validation","348","15.0%"],["Test","347","15.0%"]], [3000,3000,3360])
doc.add_heading("3.2 Character Dataset", level=2)
doc.add_paragraph("The character detector uses a Roboflow export of cropped Egyptian plates licensed under CC BY 4.0. It contains 6,332 images: 5,216 training, 745 validation, and 371 test images. Bounding boxes cover 38 classes: digits 0-9 and transliterated Arabic letter labels. The repository notes that nine base image identities are shared across supplied splits; therefore, final benchmarking must regroup samples by original plate identity to remove leakage.")

page_break(); doc.add_heading("Chapter 3: Proposed Method", level=1)
doc.add_heading("3.3 Dataset Engineering", level=2)
add_bullets(["Inspection: scan files, detect formats, compute hashes and image metadata.", "Exploratory analysis: generate dataset size, resolution, brightness, blur, class, and bounding-box plots.", "Quality control: flag corruption, missing/orphan labels, duplicates, unknown classes, and invalid boxes.", "Harmonization: parse VOC and YOLO annotations into a unified schema and normalized YOLO labels.", "Preprocessing: apply configurable transforms while preserving annotation geometry.", "Statistics and splitting: generate reports and deterministic train/validation/test manifests using fixed seeds."])
doc.add_heading("3.4 Three-Stage Recognition Cascade", level=2)
add_table(["Stage", "Input", "Model / logic", "Output"], [
    ["1. Vehicle", "Road image or frame", "COCO-pretrained YOLO11n; car, motorcycle, bus, truck", "Vehicle boxes"],
    ["2. Plate", "Padded vehicle crop", "Custom one-class YOLO11n", "Plate boxes and confidence"],
    ["3. Characters", "Plate crop", "Custom 38-class YOLO26n", "Character boxes, class, glyph, confidence"],
    ["4. Decode", "Character detections", "Row clustering; right-to-left sorting; locale map", "Formatted Arabic plate text"]], [1300,1800,3400,2860])
doc.add_paragraph("The backend maps crop-relative boxes back to the original image, combines stage confidences for suppression, and returns structured JSON. If character weights are unavailable, an optional CRNN/CTC recognizer can serve as a fallback. The application adds upload, image/video inference, training monitoring, dataset inspection, and human review feedback.")
add_figure(ROOT/"models/detection/master_plate_yolo11_20260727_132515/val_batch0_pred.jpg", "Figure 3.2. Plate-detector predictions on validation images, including vehicles and motorcycles.", 6.2)

page_break(); doc.add_heading("Chapter 3: Evaluation Metrics", level=1)
add_table(["Metric", "Definition", "Interpretation"], [
    ["IoU", "Overlap of predicted and ground-truth boxes divided by their union", "Localization agreement"],
    ["Precision", "TP / (TP + FP)", "Resistance to false detections"],
    ["Recall", "TP / (TP + FN)", "Ability to find all objects"],
    ["F1", "2PR / (P + R)", "Balance of precision and recall"],
    ["mAP@0.50", "Mean AP at IoU = 0.50", "Detection quality at a permissive overlap"],
    ["mAP@0.50:0.95", "Mean AP across IoU thresholds 0.50 to 0.95", "Stricter localization quality"],
    ["CER", "Character edit distance / reference characters", "Character transcription error"],
    ["Exact-match accuracy", "Perfect plate strings / all reference plates", "End-to-end recognition success"]], [1500,4500,3360])
doc.add_paragraph("The current repository records detector precision, recall, mAP@0.50, and mAP@0.50:0.95. Final evaluation should additionally report exact plate-string accuracy, character error rate, end-to-end recall, false positives per image, latency, and frames per second on the target hardware.")

doc.add_heading("Chapter 3: Results & Analysis (Till Now)", level=1)
add_table(["Model / run", "Best epoch", "Precision", "Recall", "mAP@0.50", "mAP@0.50:0.95"], [
    ["YOLO11 plate detector (50-epoch run)", "49", "92.02%", "91.24%", "96.24%", "82.38%"],
    ["YOLO11 plate detector (100-epoch run)", "96", "93.10%", "88.61%", "95.52%", "81.91%"],
    ["YOLO26 character detector", "27", "94.90%", "94.37%", "97.93%", "63.75%"]], [2700,1000,1300,1200,1450,1710])
doc.add_paragraph("The 50-epoch plate run is the strongest recorded checkpoint by mAP@0.50:0.95 (82.38%) and also has higher recall than the 100-epoch run. Continuing training increased precision slightly in the longer run but did not improve the stricter localization metric, suggesting diminishing returns or mild overfitting. The character detector achieves high mAP@0.50 but a materially lower mAP@0.50:0.95. It generally finds the characters, yet precise box localization remains more difficult across 38 small classes.")
doc.add_paragraph("Training curves show stable plate performance after the early epochs and continued reduction in training loss. Dataset QA also reveals the main threats to an unbiased benchmark: near duplicates in the plate sources and shared plate identities across character splits. The reported values are therefore evidence of component readiness, not a completed end-to-end recognition result.")
add_figure(ROOT/"models/detection/master_plate_yolo11_20260727_132515/results.png", "Figure 3.3. YOLO11 plate-detector training and validation curves for the 100-epoch run.", 6.35)
add_figure(ROOT/"models/character/yolo26_characters/labels.jpg", "Figure 3.4. Character-class and bounding-box distributions in the training export.", 5.9)

page_break(); doc.add_heading("Chapter 4: Conclusions and Future Work", level=1)
doc.add_heading("4.1 Conclusions", level=2)
doc.add_paragraph("The project delivers a reproducible Egyptian ALPR foundation rather than an isolated notebook. It unifies heterogeneous datasets, detects quality problems, prepares leakage-aware splits, trains plate and Arabic-character detectors, formats right-to-left output, and exposes the workflow through an API and web interface. Component validation is strong: the selected plate model reaches 96.24% mAP@0.50 and 82.38% mAP@0.50:0.95, while the character detector reaches 97.93% and 63.75%, respectively. The evidence supports continued integration testing, but it does not yet justify a final whole-plate recognition accuracy claim.")
doc.add_heading("4.2 Limitations", level=2)
add_bullets(["Character split leakage exists at the original-plate identity level.", "End-to-end exact-match accuracy, CER, latency, and video-level tracking metrics have not yet been reported.", "Vehicle detection relies on COCO pretraining rather than an Egyptian traffic-domain fine-tune.", "The character dataset is imbalanced across 38 classes and tight small-object localization remains challenging.", "Validation images include near duplicates, varied source quality, and inconsistent acquisition conditions."])
doc.add_heading("4.3 Future Work", level=2)
add_bullets(["Rebuild plate and character splits by source/vehicle/plate identity and freeze a final test set.", "Run an end-to-end benchmark with exact plate accuracy, CER, recall, false positives per image, latency, and FPS.", "Perform per-class and error-type analysis for Arabic letters, digits, two-row plates, motorcycles, blur, glare, and severe tilt.", "Compare YOLO11/YOLO26 against RF-DETR and CRNN under the same leakage-controlled protocol.", "Add temporal tracking and multi-frame voting for video.", "Use human-review corrections for active learning and targeted retraining.", "Calibrate confidence thresholds and export an optimized model for edge deployment."])

page_break(); doc.add_heading("References", level=1)
refs = [
"[1] M. El-Adawi et al., 'Automated New License Plate Recognition in Egypt,' Alexandria Engineering Journal, vol. 52, no. 3, pp. 319-326, 2013. doi: 10.1016/j.aej.2013.02.005.",
"[2] A. Abd El Rahman, A. Hamdy, and F. S. A. Zaki, 'Automatic Arabic Number Plate Recognition,' Proc. JEC-ECC, 2013. doi: 10.1109/JEC-ECC.2013.6766379.",
"[3] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, 'You Only Look Once: Unified, Real-Time Object Detection,' Proc. CVPR, 2016, pp. 779-788. https://arxiv.org/abs/1506.02640.",
"[4] S. Z. Masood, G. Shu, A. Dehghan, and E. G. Ortiz, 'License Plate Detection and Recognition Using Deeply Learned Convolutional Neural Networks,' 2017. https://arxiv.org/abs/1703.07330.",
"[5] B. Shi, X. Bai, and C. Yao, 'An End-to-End Trainable Neural Network for Image-Based Sequence Recognition and Its Application to Scene Text Recognition,' IEEE TPAMI, vol. 39, no. 11, pp. 2298-2304, 2017. https://arxiv.org/abs/1507.05717.",
"[6] A. R. Youssef, F. R. Sayed, and A. A. Ali, 'A New Benchmark Dataset for Egyptian License Plate Detection and Recognition,' Proc. ACIRS, 2022, pp. 106-111. doi: 10.1109/ACIRS55390.2022.9845514.",
"[7] A. R. Youssef, F. R. Sayed, and A. A. Ali, 'Real-time Egyptian License Plate Detection and Recognition using YOLO,' IJACSA, vol. 13, no. 7, 2022.",
"[8] M. M. Abdellatif et al., 'A Low Cost IoT-Based Arabic License Plate Recognition Model for Smart Parking Systems,' Ain Shams Engineering Journal, vol. 14, no. 6, 102178, 2023. doi: 10.1016/j.asej.2023.102178.",
"[9] 'Egyptian motorcycle and vehicle license plate recognition,' Neural Computing and Applications, 2025. doi: 10.1007/s00521-025-11302-6.",
"[10] G. Jocher and J. Qiu, 'Ultralytics YOLO11,' software, 2024. https://docs.ultralytics.com/models/yolo11/.",
"[11] Ultralytics, 'YOLO Performance Metrics,' https://docs.ultralytics.com/guides/yolo-performance-metrics/ (accessed Aug. 1, 2026).",
"[12] A. R. Youssef et al., 'Egyptian Automatic License Plate Recognition (EALPR) Dataset,' GitHub repository. https://github.com/ahmedramadan96/EALPR."]
for ref in refs:
    p=doc.add_paragraph(ref); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(6)

# Header/footer and page field
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text="EGYPTIAN AUTOMATIC LICENSE PLATE RECOGNITION"; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in hp.runs: set_font(r,size=8.5,bold=True,color=GRAY)
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=fp.add_run("Page "); set_font(r,size=9,color=GRAY)
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); fp._p.append(fld)

doc.core_properties.title = "Egyptian Automatic License Plate Recognition Using a Multi-Stage YOLO Pipeline"
doc.core_properties.subject = "Final project report"
doc.core_properties.author = "Student"
doc.save(OUT)
print(OUT)
